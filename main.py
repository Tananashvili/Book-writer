import re
import openai
from docx import Document
from docx.shared import Pt 

# INSERT OPENAI API KEY
openai.api_key = ""

def main(main_character, sex, genre, age, loc):

    response = openai.ChatCompletion.create(
        model = "gpt-3.5-turbo",
        messages = [{"role": "user", "content": 
                     f"""Please provide me with book title and the headers of each chapter for a story where the main character is {main_character.capitalize()}, a {sex}, 
                     the action takes place in the {age}, in the location of {loc}, the genre is {genre}, maximum of 5 chapters."""}])
    
    print("\n Got Chapter Names")

    reply_content = response.choices[0].message.content
    title_match = re.search(r'Title: "(.*?)"', reply_content)
    title = title_match.group(1)
    chapter_names = re.findall(r'Chapter \d+: (.+)', reply_content)

    chapters = []
    chapter_num = 1
    conclusions = []
    for chapter in chapter_names:
        conc_to_input = ' '.join(conclusions[-3:])
        content = f"""Please provide me with chapter {chapter_num}/{len(chapter_names)} of the book named {title}, you have to generate story where the main character is 
        {main_character.capitalize()}, a {sex}, the action takes place in the {age}, in the location of {loc} and the genre is {genre}. chapter {chapter_num} header is "{chapter}".
        last paragraph must contain very small conclusion of chapter and names of the characters with their roles, without header."""
        
        if chapter_num != 1:
            content += f" read what happened in the last chapters and continue story: {conc_to_input}"
        
        response = openai.ChatCompletion.create(
            model = "gpt-3.5-turbo",
            messages = [{"role": "user", "content": content}])
        
        reply_content = response.choices[0].message.content
        paragraphs = reply_content.strip().split('\n\n')
        chapter_text = paragraphs[1:-1]
        chapter_story = '\n\n'.join(chapter_text)

        if "Chapter 1 Conclusion:" in chapter_story:
            chapter_story = chapter_story.split(f"Chapter {chapter_num} Conclusion:")[0]
        else:
            chapter_story = chapter_story.split("Conclusion:")[0]
        chapters.append(chapter_story)

        conclusion = paragraphs[-1]

        print(f"Chapter {chapter_num}")
        print(f"{conclusion} \n")

        conclusions.append(conclusion)
        chapter_num += 1
    
    doc = Document()
    doc.add_heading(title, 0)
    cn = 0
    for chapter in chapters:
        chapter_header = doc.add_heading().add_run(chapter_names[cn])
        chapter_header.font.size = Pt(13)
        doc.add_paragraph(chapter)
        cn += 1
    doc.save('Book.docx')


if __name__ == "__main__":
    main_character = input("Input Main Character: ")
    sex = input("Sex of Main Character: ")
    genre = input("Genre: ")
    age = input("Which Age Action Takes Place In?: ")
    loc = input("Location: ")

    main(main_character, sex, genre, age, loc)